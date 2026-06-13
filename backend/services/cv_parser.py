from __future__ import annotations

from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF

from backend.core.schemas import CVParseResult
from backend.services.utils import normalize_text


def _env_int(name: str, default: int) -> int:
    import os

    val = os.environ.get(name)
    if not val:
        return default
    try:
        return int(val)
    except ValueError:
        return default


def _env_float(name: str, default: float) -> float:
    import os

    val = os.environ.get(name)
    if not val:
        return default
    try:
        return float(val)
    except ValueError:
        return default


def _env_str(name: str, default: str) -> str:
    import os

    return os.environ.get(name, default)


def extract_text_pymupdf(pdf_path: str | Path) -> str:
    path = Path(pdf_path)
    doc = fitz.open(path.as_posix())
    try:
        parts: list[str] = []
        for page in doc:
            parts.append(page.get_text("text"))
        return "\n".join(parts)
    finally:
        doc.close()


def extract_text_pdfplumber(pdf_path: str | Path) -> str:
    from pdfplumber import open as pdf_open

    path = Path(pdf_path)
    parts: list[str] = []
    with pdf_open(path.as_posix()) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    return "\n".join(parts)


def extract_text_ocr_pymupdf(pdf_path: str | Path) -> str:
    """
    OCR fallback for scanned PDFs (best-effort).

    Requires:
    - Python package `pytesseract`
    - Tesseract OCR installed on the machine.

    Optional:
    - Set `TESSERACT_CMD` env var to the full path of `tesseract.exe`.
    """
    try:
        import os

        import pytesseract
        from PIL import Image
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("OCR_DEPENDENCIES_MISSING") from exc

    tesseract_cmd = os.environ.get("TESSERACT_CMD")
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    # Tunables (Docker-friendly via env vars)
    dpi = _env_int("OCR_DPI", 300)
    # For Vietnamese CVs, install language pack `tesseract-ocr-vie` and set OCR_LANG="vie+eng"
    lang = _env_str("OCR_LANG", "eng")
    oem = _env_int("OCR_OEM", 1)  # LSTM
    psm = _env_int("OCR_PSM", 6)  # Assume a uniform block of text
    crop_ratio = _env_float("OCR_CROP_RATIO", 0.05)  # crop 5% margins by default

    path = Path(pdf_path)
    doc = fitz.open(path.as_posix())
    try:
        parts: list[str] = []
        for page in doc:
            # Render grayscale to reduce noise
            pix = page.get_pixmap(dpi=dpi, colorspace=fitz.csGRAY)
            img = Image.frombytes("L", [pix.width, pix.height], pix.samples)

            # Crop margins to remove black borders / scan artifacts
            if 0.0 < crop_ratio < 0.3:
                w, h = img.size
                dx = int(w * crop_ratio)
                dy = int(h * crop_ratio)
                if w - 2 * dx > 20 and h - 2 * dy > 20:
                    img = img.crop((dx, dy, w - dx, h - dy))

            config = f"--oem {oem} --psm {psm}"
            parts.append(pytesseract.image_to_string(img, lang=lang, config=config))
        return "\n".join(parts)
    finally:
        doc.close()


def extract_text_ocr_image(image_path: str | Path) -> str:
    """
    OCR for image scans (png/jpg/jpeg).
    """
    try:
        import os

        import pytesseract
        from PIL import Image
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("OCR_DEPENDENCIES_MISSING") from exc

    tesseract_cmd = os.environ.get("TESSERACT_CMD")
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    lang = _env_str("OCR_LANG", "eng")
    oem = _env_int("OCR_OEM", 1)
    psm = _env_int("OCR_PSM", 6)
    crop_ratio = _env_float("OCR_CROP_RATIO", 0.05)

    img = Image.open(Path(image_path)).convert("L")
    if 0.0 < crop_ratio < 0.3:
        w, h = img.size
        dx = int(w * crop_ratio)
        dy = int(h * crop_ratio)
        if w - 2 * dx > 20 and h - 2 * dy > 20:
            img = img.crop((dx, dy, w - dx, h - dy))

    config = f"--oem {oem} --psm {psm}"
    return pytesseract.image_to_string(img, lang=lang, config=config)


def parse_cv(pdf_path: str | Path, *, fallback_pdfplumber: bool = True) -> CVParseResult:
    path = Path(pdf_path)
    cv_id = path.name

    error: Optional[str] = None
    raw_text = ""
    method = "unknown"
    try:
        raw_text = extract_text_pymupdf(path)
        raw_text = normalize_text(raw_text)
        method = "pymupdf"

        if fallback_pdfplumber and not raw_text:
            raw_text = extract_text_pdfplumber(path)
            raw_text = normalize_text(raw_text)
            method = "pdfplumber"

        if not raw_text:
            try:
                raw_text = extract_text_ocr_pymupdf(path)
                raw_text = normalize_text(raw_text)
                method = "ocr_pdf"
            except Exception:
                raw_text = ""

        if not raw_text:
            error = "EMPTY_TEXT_NEEDS_OCR"
    except Exception as exc:  # noqa: BLE001
        error = f"PARSE_ERROR: {exc.__class__.__name__}: {exc}"

    return CVParseResult(cv_id=cv_id, raw_text=raw_text, error=error, method=method)


def extract_text_docx(docx_path: str | Path) -> str:
    import docx
    doc = docx.Document(docx_path)
    full_text: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
    return "\n".join(full_text)


def extract_text_markdown(md_path: str | Path) -> str:
    with open(md_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def parse_file(path: str | Path) -> CVParseResult:
    """
    Parse either PDF (text/OCR), image scan (OCR), Word document, or Markdown text.
    Supported: .pdf, .png, .jpg, .jpeg, .docx, .md
    """
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix in {".png", ".jpg", ".jpeg"}:
        try:
            text = normalize_text(extract_text_ocr_image(p))
            err = None if text else "EMPTY_TEXT_NEEDS_OCR"
            return CVParseResult(cv_id=p.name, raw_text=text, error=err, method="ocr_image")
        except Exception as exc:  # noqa: BLE001
            return CVParseResult(
                cv_id=p.name,
                raw_text="",
                error=f"PARSE_ERROR: {exc.__class__.__name__}: {exc}",
                method="ocr_image",
            )
    elif suffix == ".docx":
        try:
            text = normalize_text(extract_text_docx(p))
            err = None if text else "EMPTY_TEXT"
            return CVParseResult(cv_id=p.name, raw_text=text, error=err, method="docx")
        except Exception as exc:  # noqa: BLE001
            return CVParseResult(
                cv_id=p.name,
                raw_text="",
                error=f"PARSE_ERROR: {exc.__class__.__name__}: {exc}",
                method="docx",
            )
    elif suffix == ".md":
        try:
            text = normalize_text(extract_text_markdown(p))
            err = None if text else "EMPTY_TEXT"
            return CVParseResult(cv_id=p.name, raw_text=text, error=err, method="markdown")
        except Exception as exc:  # noqa: BLE001
            return CVParseResult(
                cv_id=p.name,
                raw_text="",
                error=f"PARSE_ERROR: {exc.__class__.__name__}: {exc}",
                method="markdown",
            )
    return parse_cv(p)
